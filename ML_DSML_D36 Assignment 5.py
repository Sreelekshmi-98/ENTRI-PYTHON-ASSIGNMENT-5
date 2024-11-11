from docx import Document

file_path = r"C:\Users\91980\OneDrive\Documents\ML-DSML-D36 TENSES 2 ASSIGNMENT.docx"

try:
    doc = Document(file_path)
    print("File Contents:\n")
    for paragraph in doc.paragraphs:
        print(paragraph.text)

except FileNotFoundError:
    print(f"Error: The file at '{file_path}' was not found.")
except IOError as e:
    print(f"An I/O error occurred: {e}")
except Exception as e:
    print(f"An error occurred: {e}")

    from docx import Document

def copy_docx_file(source_path, destination_path):
    try:

        source_doc = Document(source_path)

        destination_doc = Document()

        for para in source_doc.paragraphs:
            destination_doc.add_paragraph(para.text)

        destination_doc.save(destination_path)

        print(f"Contents copied from {source_path} to {destination_path} successfully.")
    except FileNotFoundError:
        print(f"Error: {source_path} not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

source_path = r"C:\Users\91980\OneDrive\Documents\ML-DSML-D36 TENSES ASSIGNMENT 1.docx"
destination_path = r"C:\Users\91980\OneDrive\Documents\ML-DSML-D36 TENSES 2 ASSIGNMENT.docx"

copy_docx_file(source_path, destination_path)

def count_words_in_file(file_path):
    try:
        with open(file_path, 'r') as file:
            content = file.read()  # Read the entire content of the file
        words = content.split()


        total_words = len(words)

        print(f"Total number of words in {file_path}: {total_words}")
    except FileNotFoundError:
        print(f"Error: The file at {file_path} was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

file_path = r"C:\Users\91980\OneDrive\Documents\ML-DSML-D36 TENSES ASSIGNMENT 1.docx"  # Replace with your file path
count_words_in_file(file_path)


user_input = input("Please enter a number: ")

try:

    number = int(user_input)
    print(f"The integer value is: {number}")
except ValueError:
    print("Oops! That was not a valid number. Please enter a valid integer.")

def check_negative(numbers):
    for num in numbers:
        if num < 0:
            raise ValueError("Negative numbers are not allowed.")

user_input = input("Please enter a list of integers separated by spaces: ")

try:

    number_list = [int(x) for x in user_input.split()]

    check_negative(number_list)

    print(f"The list of integers is: {number_list}")
except ValueError as e:
    print(f"Error: {e}")


def calculate_average(numbers):
    return sum(numbers) / len(numbers)


user_input = input("Please enter a list of integers separated by spaces: ")

try:

    number_list = [int(x) for x in user_input.split()]

    if not number_list:
        raise ValueError("The list cannot be empty.")

    average = calculate_average(number_list)

    print(f"The average of the numbers is: {average}")

except ValueError as e:

    print(f"Error: {e}")

finally:

    print("Program has finished running.")

    filename = input("Please enter the filename to write to: ")

    content = "Hello, welcome to the file writing program!"

    try:
        with open(filename, 'w') as file:
            file.write(content)
        print("Welcome! The content has been successfully written to the file.")

    except Exception as e:

        print(f"An error occurred: {e}")

    finally:
        print("Program execution finished.")



